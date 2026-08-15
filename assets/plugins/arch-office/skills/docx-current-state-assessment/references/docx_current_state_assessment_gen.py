#!/usr/bin/env python3
"""Скелет генератора оценки текущего состояния / архитектурного аудита (docx).

Запуск: python3 docx_current_state_assessment_gen.py audit.docx
"""

import sys

from docx import Document
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)

# ── ЗАПОЛНИТЬ ────────────────────────────────────────────────────────────────
TITLE = "Оценка текущего состояния: <домен/система>"
META = [
    ("Объём", "домен платежей: 6 систем, 14 стыков"),
    ("Метод", "анализ кода/конфигов, метрики, инциденты, интервью"),
    ("Период", "август 2026"),
]
RESUME = [
    "Общая оценка: работоспособно, но два узла — единые точки отказа.",
    "Топ-3 находки: …",
    "Главная рекомендация: …",
]
INVENTORY = [("Система", "стек, критичность, владелец — детально в каталоге систем")]
# Находки: (id, линза, находка, свидетельство, влияние×вероятность)
FINDINGS = [
    ("F-1", "Стойкость", "единая точка отказа на платёжном хабе",
     "k8s: replicas=1 (deploy/hub.yaml:14); инцидент INC-2211", "высокое × средняя"),
    ("F-2", "Стыки", "нет идемпотентности на приёме статусов",
     "consumer.py: обработчик без dedup (src/mq/consumer.py:88)", "высокое × высокая"),
]
# Долг: (элемент, цена поддержки/год, цена устранения, комментарий)
TECH_DEBT = [
    ("Самописный ORM-слой", "₽3.5 млн (2 FTE сопровождения)", "₽6 млн, 2 кв.", "миграция на поддерживаемый фреймворк"),
]
QUICK_WINS = ["включить autoscale хаба (2 недели, без изменений кода)"]
STRATEGIC = ["перевод статусной модели на событийную с outbox (в дорожную карту H2)"]
EVIDENCE_LOG = [("E-1", "deploy/hub.yaml:14 — replicas=1"), ("E-2", "INC-2211 — простой 40 мин")]
# ─────────────────────────────────────────────────────────────────────────────


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = NAVY
        st.font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def main(path: str) -> None:
    doc = Document()
    style_document(doc)
    doc.add_heading(TITLE, level=0)

    doc.add_heading("1. О документе", level=1)
    add_table(doc, ["Поле", "Значение"], [[k, v] for k, v in META])

    doc.add_heading("2. Резюме", level=1)
    for line in RESUME:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("3. Инвентаризация", level=1)
    add_table(doc, ["Объект", "Краткое описание"], [list(i) for i in INVENTORY])

    doc.add_heading("4. Находки", level=1)
    add_table(
        doc,
        ["ID", "Линза", "Находка", "Свидетельство", "Влияние × вероятность"],
        [list(f) for f in FINDINGS],
    )

    doc.add_heading("5. Технический долг", level=1)
    add_table(doc, ["Элемент", "Цена поддержки/год", "Цена устранения", "Комментарий"], [list(t) for t in TECH_DEBT])

    doc.add_heading("6. Рекомендации", level=1)
    doc.add_heading("6.1 Quick wins (≤ квартал)", level=2)
    for qw in QUICK_WINS:
        doc.add_paragraph(qw, style="List Bullet")
    doc.add_heading("6.2 Стратегические работы", level=2)
    for st in STRATEGIC:
        doc.add_paragraph(st, style="List Bullet")

    doc.add_heading("7. Приложение: журнал свидетельств", level=1)
    add_table(doc, ["ID", "Свидетельство"], [list(e) for e in EVIDENCE_LOG])

    doc.save(path)
    print(f"записан {path}")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "audit.docx")
