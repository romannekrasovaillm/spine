#!/usr/bin/env python3
"""Скелет генератора плана миграции/трансформации (docx).

Запуск: python3 docx_migration_roadmap_gen.py migration.docx
"""

import sys

from docx import Document
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)

# ── ЗАПОЛНИТЬ ────────────────────────────────────────────────────────────────
TITLE = "План миграции: <что> → <куда>"
META = [("Программа", "владелец, спонсор"), ("Связанные", "концепция v…, SAD v…")]
DRIVER = "Зачем переезжаем (1 абзац) и критерии успеха программы целиком."
STRATEGY = [
    "Волновая модель: каждая волна самостоятельно ценна и обратима.",
    "Strangler-фасад: новый контур забирает потоки постепенно, двойной режим ограничен по времени.",
]
# Волны: словарь по ключам ниже
WAVES = [
    {
        "name": "Волна 1 — пилот: перевод отчётных потоков (только чтение)",
        "scope": "чтения справочников и отчётов; запись остаётся в legacy",
        "deps": "—",
        "entry": "фасад развёрнут, тестовый контур нагружен профилем дня",
        "exit": "2 недели без расхождений сверки чтений; p95 в норме",
        "data": "миграция справочников, сверка 1:1 по ключам, допуск 0",
        "rollback": "переключение DNS/маршрута на legacy; данные не менялись",
    },
    {
        "name": "Волна 2 — перевод записи платежей",
        "scope": "create/confirm операций",
        "deps": "волна 1",
        "entry": "outbox на legacy, дедупликация проверена тестами",
        "exit": "доля ошибок < 0.1% за месяц; сверка с АБС без расхождений",
        "data": "обратная синхронизация статусов в legacy до отключения",
        "rollback": "флаг маршрутизации; повторная выгрузка дельты статусов",
    },
]
DATA_RECON = ["Сверка: ежедневная реконсиляция по tx_id; допуск 0; владелец — команда данных."]
CALENDAR = ["Запретные окна: отчётные даты ЦБ, новогодний пик (20.12–10.01), окно релизной заморозки."]
OPS_READINESS = ["Мониторинг нового контура до волны 2; обучение дежурной смены; runbook отката."]
RISKS = [("RP-1", "расхождение статусной модели legacy/новый", "двойная сверка, карта соответствий")]
DECOMMISSION = "Волна N — вывод legacy: критерии (0 активных потребителей 2 квартала), архивация данных, снятие с учёта."
# ─────────────────────────────────────────────────────────────────────────────


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = NAVY
        st.font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def main(path: str) -> None:
    doc = Document()
    style_document(doc)
    doc.add_heading(TITLE, level=0)

    doc.add_heading("1. О документе", level=1)
    add_table(doc, ["Поле", "Значение"], [[k, v] for k, v in META])

    doc.add_heading("2. Драйвер и критерии успеха", level=1)
    doc.add_paragraph(DRIVER)

    doc.add_heading("3. Стратегия", level=1)
    for s in STRATEGY:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("4. Волны миграции", level=1)
    for w in WAVES:
        doc.add_heading(w["name"], level=2)
        add_table(doc, ["Аспект", "Содержание"], [
            ["Состав", w["scope"]],
            ["Зависимости", w["deps"]],
            ["Критерии входа", w["entry"]],
            ["Критерии выхода", w["exit"]],
            ["Данные", w["data"]],
            ["План отката", w["rollback"]],
        ])

    doc.add_heading("5. Перенос данных и реконсиляция", level=1)
    for d in DATA_RECON:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("6. Календарные ограничения", level=1)
    for c in CALENDAR:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("7. Операционная готовность", level=1)
    for o in OPS_READINESS:
        doc.add_paragraph(o, style="List Bullet")

    doc.add_heading("8. Риски программы", level=1)
    add_table(doc, ["ID", "Риск", "Митигация"], [list(r) for r in RISKS])

    doc.add_heading("9. Вывод legacy из эксплуатации", level=1)
    doc.add_paragraph(DECOMMISSION)

    doc.save(path)
    print(f"записан {path}")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "migration.docx")
