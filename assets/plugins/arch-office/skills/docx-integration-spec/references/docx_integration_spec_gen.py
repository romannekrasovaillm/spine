#!/usr/bin/env python3
"""Скелет генератора интеграционной спецификации (docx).

Запуск: python3 docx_integration_spec_gen.py spec.docx
"""

import sys

from docx import Document
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)

# ── ЗАПОЛНИТЬ ────────────────────────────────────────────────────────────────
TITLE = "Интеграционная спецификация: <система A> ↔ <система B>"
META = [
    ("Поставщик", "Система A (команда, владелец)"),
    ("Потребитель", "Система B (команда, владелец)"),
    ("Версия контракта", "1.0 (статус: проект)"),
]
FLOW_ASCII = "Потребитель        Поставщик\n     │── запрос ──────▶│\n     │◀──── ответ ─────│"
TRANSPORT = [
    ("Протокол", "HTTPS/REST JSON (или AMQP 0-9-1 и т.п.)"),
    ("Аутентификация", "mTLS + токен; ротация сертификатов — ежегодно"),
    ("Каналы", "промышленный / тестовый (стенд): URL…"),
]
# Операции: (название, семантика, пример сообщения, [(код, смысл, поведение вызывающего)])
OPERATIONS = [
    (
        "POST /payments — создать платёж",
        "Создаёт платёж; при повторе с тем же X-Idempotency-Key возвращает ранее созданный.",
        '{"tx_id": "a1b2", "amount": "100.00", "currency": "RUB"}',
        [
            ("200", "принято в обработку", "—"),
            ("409", "конфликт ключа идемпотентности", "запросить статус по tx_id, не повторять"),
            ("5xx/таймаут", "сбой поставщика", "ретрай по политике раздела 6"),
        ],
    ),
]
IDEMPOTENCY = [
    ("Ключ", "заголовок X-Idempotency-Key (UUID), обязателен для POST"),
    ("Окно хранения", "72 часа; повтор в окне → тот же ответ, повтор после — 422"),
    ("Дедупликация событий", "по event_id, окно 24 часа"),
]
RETRIES = [
    ("Таймаут ответа", "5 с connect / 30 с read"),
    ("Ретраи", "до 3 попыток, backoff 1с→4с→16с + джиттер ±20%, только на 5xx/таймаут/reset"),
    ("Общий дедлайн операции", "60 с, далее — разбор по статусной модели (раздел 4)"),
]
SLA = [("p95 ответа", "≤ 800 мс (окно: рабочий день, источник: APM)"), ("RPS", "до 50; burst 100 на 30 с")]
VERSIONING = [
    "Совместимость: добавлять можно только необязательные поля; удаление/переименование — новая major-версия.",
    "Sunset: поддержка N−1 версии 6 месяцев с анонсом.",
]
TESTS = [
    ("Позитив", "валидный запрос → 200, платёж в статусе NEW"),
    ("Дубль", "повтор с тем же ключом → тот же 200 без нового платежа"),
    ("Таймаут", "поставщик молчит 30 с → ретрай-политика потребителя"),
    ("Невалидное сообщение", "400 без побочных эффектов"),
]
# ─────────────────────────────────────────────────────────────────────────────


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = NAVY
        st.font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def add_mono(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)


def main(path: str) -> None:
    doc = Document()
    style_document(doc)
    doc.add_heading(TITLE, level=0)

    doc.add_heading("1. О документе", level=1)
    add_table(doc, ["Поле", "Значение"], [[k, v] for k, v in META])

    doc.add_heading("2. Схема взаимодействия", level=1)
    add_mono(doc, FLOW_ASCII)

    doc.add_heading("3. Транспорт и безопасность", level=1)
    add_table(doc, ["Параметр", "Значение"], [list(t) for t in TRANSPORT])

    doc.add_heading("4. Операции", level=1)
    for name, semantics, example, errors in OPERATIONS:
        doc.add_heading(name, level=2)
        doc.add_paragraph(semantics)
        add_mono(doc, example)
        add_table(doc, ["Код", "Смысл", "Поведение вызывающего"], [list(e) for e in errors])

    doc.add_heading("5. Идемпотентность и дедупликация", level=1)
    add_table(doc, ["Механизм", "Правило"], [list(i) for i in IDEMPOTENCY])

    doc.add_heading("6. Таймауты и ретраи", level=1)
    add_table(doc, ["Параметр", "Значение"], [list(r) for r in RETRIES])

    doc.add_heading("7. SLA и лимиты", level=1)
    add_table(doc, ["Метрика", "Цель"], [list(s) for s in SLA])

    doc.add_heading("8. Версионирование", level=1)
    for v in VERSIONING:
        doc.add_paragraph(v, style="List Bullet")

    doc.add_heading("9. Тестовые сценарии", level=1)
    add_table(doc, ["Сценарий", "Ожидаемый результат"], [list(t) for t in TESTS])

    doc.save(path)
    print(f"записан {path}")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "spec.docx")
