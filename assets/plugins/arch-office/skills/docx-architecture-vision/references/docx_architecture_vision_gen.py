#!/usr/bin/env python3
"""Скелет генератора архитектурной концепции/vision (docx, для правления).

Запуск: python3 docx_architecture_vision_gen.py vision.docx
"""

import sys

from docx import Document
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)

# ── ЗАПОЛНИТЬ ────────────────────────────────────────────────────────────────
TITLE = "Архитектурная концепция: <направление> 2026–2029"
META = [("Статус", "Проект"), ("Владелец", "Команда архитектуры"), ("Связанные", "Аналитический отчёт №…")]
BOARD_RESUME = [
    "Цель: …",
    "Эффект: ₽X млн/год (допущения — раздел 8).",
    "Запрашиваемое решение: утвердить горизонт H1 и финансирование …",
]
DRIVERS = ["Рынок: …", "Регуляторика: …", "Технологии: …"]
CURRENT = ["Что работает сегодня и в чём его предел (цифры, не эпитеты)."]
TARGET_CAPABILITIES = ["Возможность 1 (без вендорских названий).", "Возможность 2…"]
PRINCIPLES = [("P-1", "Принцип — проверяемая формулировка и как контролируется")]
GAPS = [("Разрыв 1", "работа по закрытию, горизонт")]
# Дорожная карта: (горизонт, содержание, результат, критерий перехода)
ROADMAP = [
    ("H1", "пилот контуров A", "работающий пилот на 2 доменах", "метрики пилота в норме 2 месяца"),
    ("H2", "промышленное распространение", "50% потока на целевом контуре", "SLA 2 квартала"),
    ("H3", "полный перевод", "вывод legacy", "отключение последнего узла"),
]
# Экономика: (статья, оценка, допущения)
ECONOMICS = [("Эффект на ФОТ", "−₽X млн/год", "база: 40 FTE × средняя ставка")]
RISKS_CHANGE = [("Риск перемен", "митигация")]
RISKS_INACTION = [("Риск бездействия", "оценка последствий")]
ASK = "Запрашиваемое решение одним абзацем: что утвердить, что профинансировать, что поручить."
# ─────────────────────────────────────────────────────────────────────────────


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = NAVY
        st.font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def main(path: str) -> None:
    doc = Document()
    style_document(doc)
    doc.add_heading(TITLE, level=0)

    doc.add_heading("1. О документе", level=1)
    add_table(doc, ["Поле", "Значение"], [[k, v] for k, v in META])

    doc.add_heading("2. Резюме для правления", level=1)
    for line in BOARD_RESUME:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("3. Драйверы перемен", level=1)
    for d in DRIVERS:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("4. Текущее состояние и его пределы", level=1)
    for c in CURRENT:
        doc.add_paragraph(c)

    doc.add_heading("5. Целевое состояние", level=1)
    doc.add_heading("5.1 Целевые возможности", level=2)
    for cap in TARGET_CAPABILITIES:
        doc.add_paragraph(cap, style="List Bullet")
    doc.add_heading("5.2 Принципы", level=2)
    add_table(doc, ["ID", "Принцип и контроль"], [list(p) for p in PRINCIPLES])

    doc.add_heading("6. Разрывы и работы", level=1)
    add_table(doc, ["Разрыв", "Работа по закрытию"], [list(g) for g in GAPS])

    doc.add_heading("7. Дорожная карта", level=1)
    add_table(doc, ["Горизонт", "Содержание", "Результат", "Критерий перехода"], [list(r) for r in ROADMAP])

    doc.add_heading("8. Экономика эффекта", level=1)
    add_table(doc, ["Статья", "Оценка", "Допущения"], [list(e) for e in ECONOMICS])

    doc.add_heading("9. Риски", level=1)
    doc.add_heading("9.1 Риски перемен", level=2)
    add_table(doc, ["Риск", "Митигация"], [list(r) for r in RISKS_CHANGE])
    doc.add_heading("9.2 Риски бездействия", level=2)
    add_table(doc, ["Риск", "Последствия"], [list(r) for r in RISKS_INACTION])

    doc.add_heading("10. Запрашиваемое решение", level=1)
    doc.add_paragraph(ASK)

    doc.save(path)
    print(f"записан {path}")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "vision.docx")
