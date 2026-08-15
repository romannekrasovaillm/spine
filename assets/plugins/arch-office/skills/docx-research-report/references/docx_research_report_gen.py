#!/usr/bin/env python3
"""Скелет генератора аналитического отчёта для МД (docx, домовой стиль отчётов для МД).

Заполнение: TITLE/SUBJECT/RESUME/SECTIONS/MATRIX*/SOURCES — и запуск:
    python3 docx_research_report_gen.py out.docx
Проверка: открыть и прочитать заголовки (python-docx) — см. чек-лист в SKILL.md.
"""

import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)   # корпоративный тёмно-синий (заголовки)
GRAY = RGBColor(0x59, 0x59, 0x59)   # служебный текст

# ── ЗАПОЛНИТЬ ────────────────────────────────────────────────────────────────
TITLE = "Аналитический отчёт: <предмет>"
SUBJECT = "Предмет, объём и метод оценки (1–2 абзаца)."
RESUME = [
    "Главный вывод одним предложением.",
    "Второй ключевой факт/цифра.",
    "Рекомендация (что предлагается сделать).",
]
# Разделы: (заголовок, [абзацы])
SECTIONS = [
    ("Ландшафт", ["Обзор предметной области: игроки, классы решений, динамика."]),
    ("Сильные стороны и ограничения", ["По каждому рассматриваемому объекту — оба списка."]),
]
# Сравнительная матрица: критерии × объекты (шкала 1–5).
MATRIX_LEGEND = "Шкала: 1 — слабо, 5 — отлично. Оценки на дату источников."
MATRIX_HEADERS = ["Критерий", "Вариант A", "Вариант B", "Вариант C"]
MATRIX_ROWS = [
    ["Зрелость", "4", "3", "2"],
    ["Стоимость владения", "3", "4", "3"],
]
RECOMMENDATIONS = [
    "Рекомендация 1 — сформулирована как решение, а не наблюдение.",
    "Рекомендация 2 — с владельцем и горизонтом.",
]
CONCLUSIONS = ["Выводы: 3–5 тезисов, каждый подкреплён телом отчёта."]
SOURCES = [("Название источника", "https://example.org", "2026-08-15")]
# ─────────────────────────────────────────────────────────────────────────────


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = NAVY
        st.font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def main(path: str) -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_heading(TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("О документе", level=1)
    p = doc.add_paragraph(SUBJECT)
    p.runs[0].font.color.rgb = GRAY

    doc.add_heading("Резюме", level=1)
    for line in RESUME:
        doc.add_paragraph(line, style="List Bullet")

    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            doc.add_paragraph(text)

    doc.add_heading("Сравнительная оценка", level=1)
    doc.add_paragraph(MATRIX_LEGEND).runs[0].font.italic = True
    add_table(doc, MATRIX_HEADERS, MATRIX_ROWS)

    doc.add_heading("Рекомендации", level=1)
    for rec in RECOMMENDATIONS:
        doc.add_paragraph(rec, style="List Number")

    doc.add_heading("Выводы", level=1)
    for con in CONCLUSIONS:
        doc.add_paragraph(con)

    doc.add_heading("Источники", level=1)
    for name, url, date in SOURCES:
        doc.add_paragraph(f"{name} — {url} (обращение: {date})", style="List Bullet")

    doc.save(path)
    print(f"записан {path}")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "report.docx")
