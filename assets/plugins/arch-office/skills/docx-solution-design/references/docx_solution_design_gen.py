#!/usr/bin/env python3
"""Скелет генератора документа архитектурного решения (SAD, docx).

Заполнение структур ниже и запуск:
    python3 docx_solution_design_gen.py sad.docx
"""

import sys

from docx import Document
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)
GRAY = RGBColor(0x59, 0x59, 0x59)

# ── ЗАПОЛНИТЬ ────────────────────────────────────────────────────────────────
TITLE = "Архитектурное решение: <название инициативы>"
META = [("Статус", "На ревью"), ("Версия", "0.1"), ("Автор", "Команда архитектуры"), ("ADR", "ADR-014, ADR-017")]
DRIVERS = [
    "Бизнес-цель и измеримый результат.",
    "Ограничение (регуляторика/стек/сроки/бюджет).",
]
FUNCTIONAL = ["FR-1: система делает …", "FR-2: пользователь может …"]
# NFR: (id, требование, цель, метод проверки)
NFRS = [
    ("NFR-1", "Доступность", "99.9% в окне 24/7", "синтетические пробы + отчёт мониторинга"),
    ("NFR-2", "Время отклика p95", "≤ 300 мс на операцию", "нагрузочный тест профилем прод-дня"),
]
# Варианты: (название, суть, плюсы, минусы/следствия)
OPTIONS = [
    ("A — <вариант>", "краткая суть", "плюсы", "минусы и цена"),
    ("B — <вариант>", "краткая суть", "плюсы", "минусы и цена"),
    ("0 — ничего не делать", "статус-кво", "нулевые затраты", "проблема остаётся и растёт"),
]
CHOICE_RATIONALE = "Почему выбран вариант X: трассировка к драйверам и NFR."
# Диаграмма — вывод инструмента mermaid_render (вставляется моноширинным блоком).
DIAGRAM_ASCII = "┌────────┐   ┌────────┐\n│ Клиент │──▶│  API   │\n└────────┘   └────────┘"
INTERFACES = [("Канал → API Gateway", "REST/JSON, см. интеграционную спецификацию IS-12")]
RISKS = [("R-1", "риск и его митигация; детали — в реестре рисков")]
ROLLOUT = ["Этап 1: … (критерий выхода: …)", "Этап 2: …"]
ROLLBACK = "План отката: точки возврата, данные, команда, критерии запуска."
# ─────────────────────────────────────────────────────────────────────────────


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = NAVY
        st.font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def main(path: str) -> None:
    doc = Document()
    style_document(doc)
    doc.add_heading(TITLE, level=0)

    doc.add_heading("1. О документе", level=1)
    add_table(doc, ["Поле", "Значение"], [[k, v] for k, v in META])

    doc.add_heading("2. Контекст и драйверы", level=1)
    for d in DRIVERS:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("3. Требования", level=1)
    doc.add_heading("3.1 Функциональные", level=2)
    for f in FUNCTIONAL:
        doc.add_paragraph(f, style="List Bullet")
    doc.add_heading("3.2 Нефункциональные (NFR)", level=2)
    add_table(doc, ["ID", "Требование", "Цель", "Проверка"], [list(n) for n in NFRS])

    doc.add_heading("4. Варианты решения", level=1)
    add_table(doc, ["Вариант", "Суть", "Плюсы", "Минусы/следствия"], [list(o) for o in OPTIONS])

    doc.add_heading("5. Выбранное решение", level=1)
    doc.add_paragraph(CHOICE_RATIONALE)
    doc.add_heading("5.1 Диаграмма (C4-контейнеры)", level=2)
    for line in DIAGRAM_ASCII.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"   # моноширинный: арт не расползается
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)

    doc.add_heading("6. Стыки и контракты", level=1)
    add_table(doc, ["Стык", "Контракт"], [list(i) for i in INTERFACES])

    doc.add_heading("7. Риски и митигации", level=1)
    add_table(doc, ["ID", "Риск и митигация"], [list(r) for r in RISKS])

    doc.add_heading("8. План внедрения и отката", level=1)
    for step in ROLLOUT:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph(ROLLBACK)

    doc.save(path)
    print(f"записан {path}")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "sad.docx")
